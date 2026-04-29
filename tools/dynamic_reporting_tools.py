"""
PyDynamicReporting 工具：通过 Ansys PyDynamicReporting 自动生成电机仿真分析报告。
支持插入文本、数据表格、图片，最终导出为 HTML 或 PDF。
每个函数返回包含 'success'、'result' 和可选 'error' 字段的字典。
"""

from __future__ import annotations

import os
from tools.utils import _ok, _err, append_warnings

_report_session = None   # 全局报告会话
_report_items: list[dict] = []   # 报告内容项列表（备用：无 ADR 时用 HTML 模板）


# ---------------------------------------------------------------------------
# 工具：create_report_session - 初始化报告会话
# ---------------------------------------------------------------------------

def create_report_session(
    title: str = "电机仿真分析报告",
    output_dir: str = "",
    use_adr: bool = True,
) -> dict:
    """
    初始化报告生成会话。

    Args:
        title: 报告标题
        output_dir: 报告输出目录；为空则使用当前工作目录
        use_adr: True 则尝试使用 Ansys Dynamic Reporting（ADR）；
                 False 则回退到内置 HTML 模板生成器
    """
    global _report_session, _report_items
    _report_items = []
    warnings: list[str] = []
    output_dir = output_dir or os.getcwd()
    if not title.strip():
        return _err("title 不能为空")
    os.makedirs(output_dir, exist_ok=True)

    if use_adr:
        try:
            import ansys.dynamicreporting.core as adr
            _report_session = adr.Service(
                ansys_installation=os.environ.get("ANSYS_INSTALLATION", ""),
                db_directory=os.path.join(output_dir, "adr_db"),
            )
            _report_session.connect(create_db=True)
            _report_session.renderer.title = title
            return _ok({
                "title": title,
                "output_dir": output_dir,
                "backend": "Ansys Dynamic Reporting",
            })
        except Exception as e:
            # ADR 不可用，回退到内置HTML模板
            _report_session = None
            warnings.append(f"ADR 不可用，已回退到内置 HTML 模板: {e}")

    # 内置 HTML 模板模式
    _report_session = {
        "type": "html",
        "title": title,
        "output_dir": output_dir,
        "sections": [],
    }
    return _ok(append_warnings({
        "title": title,
        "output_dir": output_dir,
        "backend": "内置 HTML 模板（ADR 不可用）",
    }, warnings))


def _check_session():
    if _report_session is None:
        raise RuntimeError("报告会话未初始化，请先调用 create_report_session。")
    return _report_session


# ---------------------------------------------------------------------------
# 工具：add_report_section - 添加文本节
# ---------------------------------------------------------------------------

def add_report_section(
    title: str,
    content: str,
    level: int = 2,
) -> dict:
    """
    向报告中添加一个文本节（段落或说明性文字）。

    Args:
        title: 小节标题
        content: 正文内容（支持 Markdown 格式的加粗、换行等）
        level: 标题级别，2=H2，3=H3
    """
    try:
        session = _check_session()
        if not title.strip():
            return _err("title 不能为空")
        if not content.strip():
            return _err("content 不能为空")
        if isinstance(session, dict) and session.get("type") == "html":
            session["sections"].append({
                "type": "section",
                "title": title,
                "content": content,
                "level": level,
            })
        else:
            # ADR 模式
            item = session.create_item(content=content, type="string")
            item.set_tags(f"section_title={title}")
            item.visualizations = item.DEFAULT_VISUALIZATIONS

        _report_items.append({"type": "section", "title": title})
        return _ok({
            "title": title,
            "type": "section",
            "message": f"已添加文本节：{title}",
        })
    except Exception as e:
        return _err(str(e))


# ---------------------------------------------------------------------------
# 工具：add_table_to_report - 添加数据表格
# ---------------------------------------------------------------------------

def add_table_to_report(
    data: list[dict],
    table_title: str = "数据表格",
) -> dict:
    """
    向报告中插入数据表格，data 为字典列表（每个字典代表一行，key 为列名）。

    Args:
        data: 表格数据，格式为 [{"列1": 值1, "列2": 值2, ...}, ...]
        table_title: 表格标题
    """
    try:
        session = _check_session()
        if not data:
            return _err("data 不能为空列表")
        if not all(isinstance(row, dict) for row in data):
            return _err("data 必须为字典列表")

        if isinstance(session, dict) and session.get("type") == "html":
            session["sections"].append({
                "type": "table",
                "title": table_title,
                "data": data,
            })
        else:
            # ADR 模式：用 table item
            import pandas as pd
            df = pd.DataFrame(data)
            item = session.create_item(content=df.to_csv(index=False), type="table")
            item.set_tags(f"table_title={table_title}")

        _report_items.append({"type": "table", "title": table_title, "rows": len(data)})
        return _ok({
            "title": table_title,
            "type": "table",
            "rows": len(data),
            "message": f"已添加表格：{table_title}（{len(data)} 行）",
        })
    except Exception as e:
        return _err(str(e))


# ---------------------------------------------------------------------------
# 工具：add_image_to_report - 添加图片/云图
# ---------------------------------------------------------------------------

def add_image_to_report(
    image_path: str,
    caption: str = "",
    width_pct: int = 80,
) -> dict:
    """
    向报告中插入图片（仿真云图、效率 MAP 截图等）。

    Args:
        image_path: 图片文件绝对路径（支持 PNG/JPG/SVG）
        caption: 图片说明文字
        width_pct: 图片在页面中的宽度百分比（1~100）
    """
    try:
        session = _check_session()
        if not os.path.exists(image_path):
            return _err(f"图片文件不存在：{image_path}")
        if width_pct <= 0 or width_pct > 100:
            return _err("width_pct 必须在 1 到 100 之间")

        if isinstance(session, dict) and session.get("type") == "html":
            session["sections"].append({
                "type": "image",
                "path": image_path,
                "caption": caption,
                "width_pct": width_pct,
            })
        else:
            # ADR 模式
            item = session.create_item(content=image_path, type="image")
            item.set_tags(f"caption={caption}")

        _report_items.append({"type": "image", "path": image_path, "caption": caption})
        return _ok({
            "type": "image",
            "path": image_path,
            "caption": caption,
            "message": f"已添加图片：{os.path.basename(image_path)}（{caption}）",
        })
    except Exception as e:
        return _err(str(e))


# ---------------------------------------------------------------------------
# 工具：export_report - 导出报告
# ---------------------------------------------------------------------------

def export_report(
    format: str = "html",
    filename: str = "motor_analysis_report",
) -> dict:
    """
    将当前报告导出为 HTML、PDF 或 Word (.docx) 文件。

    Args:
        format: "html"、"pdf" 或 "docx"（PDF 需要 ADR 支持，
                docx 基于 python-docx 生成符合工程报告规范的 Word 文档）
        filename: 输出文件名（不含扩展名）
    """
    try:
        session = _check_session()
        if not filename.strip():
            return _err("filename 不能为空")

        if isinstance(session, dict) and session.get("type") == "html":
            if format == "pdf":
                return _err("内置 HTML 模板后端不支持 PDF 导出，请启用 ADR 或安装 PDF 渲染后端")
            if format not in ("html", "docx"):
                return _err(f"不支持的导出格式: {format}；内置模板后端仅支持 html 或 docx")
        if not _report_items:
            return _err("当前报告内容为空，请先添加节、表格或图片后再导出")

        if isinstance(session, dict) and session.get("type") == "html":
            if format == "docx":
                output_path = os.path.join(session["output_dir"], f"{filename}.docx")
                _render_docx_report(session, output_path)
            else:
                output_path = os.path.join(session["output_dir"], f"{filename}.html")
                _render_html_report(session, output_path)
        else:
            # ADR 导出
            output_path = os.path.join(
                session.renderer.report_directory, f"{filename}.{format}"
            )
            if format == "pdf":
                session.renderer.render(output_path, format="pdf")
            elif format == "docx":
                # ADR 不直接支持 docx，回退到内置 docx 渲染器
                _render_docx_report(session, output_path)
            else:
                session.renderer.render(output_path)
        if not os.path.exists(output_path):
            return _err(f"报告导出后未生成文件: {output_path}")

        return _ok({
            "output_path": output_path,
            "format": format,
            "num_sections": len(_report_items),
        })
    except Exception as e:
        return _err(str(e))


def _render_html_report(session: dict, output_path: str) -> None:
    """内置 HTML 模板渲染器。"""
    title = session["title"]
    sections = session.get("sections", [])

    css = """
    body { font-family: Arial, sans-serif; max-width: 1200px; margin: 40px auto; padding: 0 20px; }
    h1 { color: #003366; border-bottom: 2px solid #003366; }
    h2 { color: #0055aa; margin-top: 30px; }
    h3 { color: #0077cc; }
    table { border-collapse: collapse; width: 100%; margin: 16px 0; }
    th { background: #003366; color: white; padding: 8px 12px; }
    td { border: 1px solid #ccc; padding: 6px 12px; }
    tr:nth-child(even) { background: #f5f5f5; }
    img { display: block; margin: 12px auto; border: 1px solid #ddd; }
    .caption { text-align: center; color: #666; font-size: 0.9em; margin-bottom: 16px; }
    """

    body_parts = [f"<h1>{title}</h1>"]
    for sec in sections:
        sec_type = sec.get("type")
        if sec_type == "section":
            lvl = sec.get("level", 2)
            body_parts.append(f"<h{lvl}>{sec['title']}</h{lvl}>")
            body_parts.append(f"<p>{sec['content'].replace(chr(10), '<br>')}</p>")
        elif sec_type == "table":
            body_parts.append(f"<h3>{sec['title']}</h3>")
            data = sec.get("data", [])
            if data:
                cols = list(data[0].keys())
                header = "".join(f"<th>{c}</th>" for c in cols)
                rows = ""
                for row in data:
                    cells = "".join(f"<td>{row.get(c, '')}</td>" for c in cols)
                    rows += f"<tr>{cells}</tr>"
                body_parts.append(f"<table><thead><tr>{header}</tr></thead><tbody>{rows}</tbody></table>")
        elif sec_type == "image":
            w = sec.get("width_pct", 80)
            caption = sec.get("caption", "")
            img_path = sec["path"].replace("\\", "/")
            body_parts.append(f'<img src="{img_path}" style="width:{w}%">')
            if caption:
                body_parts.append(f'<p class="caption">{caption}</p>')

    html = f"<!DOCTYPE html><html><head><meta charset='utf-8'><title>{title}</title><style>{css}</style></head><body>{''.join(body_parts)}</body></html>"
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)


def _render_docx_report(session: dict, output_path: str) -> None:
    """基于 python-docx 生成符合工程报告规范的 Word 文档。"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise RuntimeError(
            "python-docx 未安装，请执行: pip install python-docx>=1.1.0"
        )

    title = session["title"]
    sections = session.get("sections", [])

    doc = Document()

    # ── 全局样式 ──────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    # ── 封面标题 ──────────────────────────────────────────────
    p_title = doc.add_heading(title, level=0)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # 空行分隔

    # ── 目录占位（Word 打开后可手动更新域） ─────────────────
    p_toc = doc.add_paragraph()
    p_toc.add_run("[ 目录 — 在 Word 中按 Ctrl+A → F9 更新域 ]").italic = True
    doc.add_page_break()

    # ── 正文各节 ──────────────────────────────────────────────
    for sec in sections:
        sec_type = sec.get("type")

        if sec_type == "section":
            level = min(max(sec.get("level", 2), 1), 4)
            doc.add_heading(sec["title"], level=level)
            # 保留换行，逐段落写入
            for para in sec["content"].split("\n"):
                para = para.strip()
                if para:
                    doc.add_paragraph(para)

        elif sec_type == "table":
            doc.add_heading(sec.get("title", "数据表格"), level=3)
            data = sec.get("data", [])
            if not data:
                doc.add_paragraph("（无数据）")
                continue
            cols = list(data[0].keys())
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 表头
            for ci, col_name in enumerate(cols):
                cell = table.rows[0].cells[ci]
                cell.text = str(col_name)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            # 数据行
            for row_data in data:
                row_cells = table.add_row().cells
                for ci, col_name in enumerate(cols):
                    row_cells[ci].text = str(row_data.get(col_name, ""))

        elif sec_type == "image":
            img_path = sec["path"]
            caption = sec.get("caption", "")
            width_pct = sec.get("width_pct", 80)
            if os.path.exists(img_path):
                # 将百分比转为英寸（A4 宽约 6.5 英寸）
                width_inch = 6.5 * width_pct / 100
                doc.add_picture(img_path, width=Inches(width_inch))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_cap.add_run(caption)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── 页脚信息 ──────────────────────────────────────────────
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run("本报告由 AnsysAgent 自动生成 | Powered by Ansys + DeepSeek")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(output_path)
